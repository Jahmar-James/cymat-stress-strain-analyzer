from docx import Document

class WordExporter:
    def __init__(self, app):
        self.app = app
        self.doc = Document()    

    def add_heading(self, text, level=0):
        self.doc.add_heading(text, level=level)

    def add_section(self, section_id, section_text):
        self.add_heading(section_text, level=1)
        # add logic here to add data from properties_df to the document
        if section_id in self.properties_df.columns:
            data_text = str(self.properties_df[section_id].values)
            self.doc.add_paragraph(data_text)
        else:
            self.doc.add_paragraph('...')  # Placeholder
   
    def df_to_docx_table(self, doc, df):
        table = doc.add_table(rows=1, cols=len(df.columns))  # create table with header row
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = str(column)  # populate header row

        for i in range(len(df)):  # populate table rows
            cells = table.add_row().cells
            for j, column in enumerate(df.columns):
                cells[j].text = str(df.iloc[i, j])

        return doc


    def export_report(self, selected_indices, file_path):
        self.add_heading('Test Report', 0)

        sections = {
        'a': 'Nature and designation of the product being tested',
        'b': 'Material and density of the sample',
        'c': 'Type of sampling',
        'd': 'Dimensions of the samples in mm, giving the ratio of side length to height',
        'e': 'Number of samples',
        'f': 'Test conditions (temperature, humidity)',
        'g': 'Type of testing machines and measuring range of the measuring equipment',
        'h': 'Compression speed(s) in s-1',
        'i': 'Plateau stress Rplt',
        'j': 'Compression at plateau end AAt / start of compression range',
        'k': 'Gradient of the elastic straight line m',
        'l': 'Energy absorption Ev',
        'm': 'Energy absorption efficiency Eff',
        'n': 'Upper compressive yield strength ReH',
        'o': 'Offset yield strength Rp1',
        's': 'a stress/strain diagram with characteristic values in each case',
        't': 'any deviations from this standard',
        'u': 'the test date'
        }   
        self.app.data_handler.update_properties_df(selected_indices)
        self.properties_df =  self.app.data_handler.properties_df

        for key in sections:
            self.add_section(key, sections[key])


        self.save(file_path)

    def save(self,file_path):
        self.doc.save(file_path)

