from docx import Document
from datetime import date
import pandas as pd

"""
DIN 51233  Material testing machines Safety specifications General specifications 
DIN 50106  Testing of metallic materials Compression test
DIN EN ISO 7500-1 Metallic materials Verification of static uniaxial testing machines Part 1: Tension- and compression testing machines Verification and calibration of the force-measuring system
DIN EN ISO 9513 Metallic materials Calibration of extensometers used in uniaxial testing 
"""

sections = {
        '1. Test Details': {
            'a': {
                'title': 'Nature and designation of the product being tested',
                'content': '[Product Details]'
            },
            'b': {
                'title': 'Material and density of the sample',
                'content': '[Material Details]'
            },
            'c': {
                'title': 'Type of sampling',
                'content': '[Sampling Details, df'
            }
        },
        '2. Sample Details': {
            'a': {
                'title': 'Dimensions of the samples in mm, giving the ratio of side length to height',
                'content': '[Sample Dimensions]'
            },
            'b': {
                'title': 'Number of samples',
                'content': '[Sample Count]'
            }
        },
        '3. Test Conditions': {
            'a': {
                'title': 'Temperature',
                'content': '[Temperature Details]'
            },
            'b': {
                'title': 'Humidity',
                'content': '[Humidity Details]'
            }
        },
        '4. Testing Machines': {
            'a': {
                'title': 'Type of testing machines',
                'content': '[Machine Details]'
            },
            'b': {
                'title': 'Measuring range of the equipment',
                'content': '[Range Details]'
            }
        },
        '5. Compression Speed(s) in s-1': {
            'a': {
                'title': 'Compression Speed Details',
                'content': '[Compression Speed Details]'
            }
        },
        '6. Evaluation Results': {
            'a': {
                'title': 'Plateau stress Rplt',
                'content': 'There is the property {Rplt}, it represents the avrage stress over the plateu region 0.2 - 0.4, the mean value for these samples is {mean Rplt Value]} '
            },
            'b': {
                'title': 'Compression at plateau end AAt / start of compression range',
                'content': '[AAt Value]'
            },
            'c': {
                'title': 'Gradient of the elastic straight line m',
                'content': '[m Value]'
            },
            'd': {
                'title': 'Energy absorption Ev',
                'content': '[Ev Value]'
            },
            'e': {
                'title': 'Energy absorption efficiency Eff',
                'content': '[Eff Value]'
            },
            'f': {
                'title': 'Upper compressive yield strength ReH',
                'content': '[ReH Value]'
            },
            'g': {
                'title': 'Offset yield strength Rp1',
                'content': '[Rp1 Value]'
            },
            'h': {
                'title': 'Stress/strain diagram with characteristic values',
                'content': '[Attach Diagram]'
            }
        },
        '7. Deviations': {
            'a': {
                'title': 'Any deviations from this standard',
                'content': '[Deviations Details]'
            }
        }
    }  

class WordExporter:
    def __init__(self, app):
        self.app = app
        self.doc = Document()
        self.Title = "DIN Compression Testing Report"
        self.date = f"Date: {str(date.today())}"
        self.standard_num =  [" DIN 51233"]
        

    def add_heading(self, text, level=0):
        self.doc.add_heading(text, level=level)
        
    def add_section(self, section_name, section_dict):
        self.add_heading(section_name, level=1)
        for subsection in section_dict.values():
            self.add_subsection(subsection['title'], subsection['content'])
            
    def add_subsection(self, title, content):
        self.add_heading(title, level=2)
        self.doc.add_paragraph()  # Add a line break after the title
        if isinstance(content, list):
            self.doc.add_paragraph(content[0])  
            if isinstance(content[1], pd.DataFrame):  
                self.df_to_docx_table(self.doc, content[1])  # Add DataFrame as a table
        elif isinstance(content, str):
            self.doc.add_paragraph(content)
  
    def df_to_docx_table(self, doc, df):
        table = doc.add_table(rows=1, cols=len(df.columns))  # create table with header row
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = str(column)  # populate header row

        for i in range(len(df)):  # populate table rows
            cells = table.add_row().cells
            for j, column in enumerate(df.columns):
                cells[j].text = str(df.iloc[i, j])

        return doc
    
    def fetch_data(self,selected_indices):
        self.app.data_handler.update_properties_df(selected_indices)
        self.properties_df =  self.app.data_handler.properties_df
        self.summary_df = self.app.data_handler.summary_statistics()
        self.mean_df = self.properties_df.mean(column= 0)
        

    def export_report(self, selected_indices, file_path):
        self.fetch_data(selected_indices)
        self.add_heading(self.Title, 0)
        self.doc.add_paragraph(self.date)
        self.doc.add_paragraph(self.standard_num)

    
        for section_name, section_dict in sections.items():
            self.add_section(section_name, section_dict)

        self.save(file_path)

    def save(self,file_path):
        self.doc.save(file_path)

